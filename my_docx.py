from docx import Document
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE


class MyDocx():
    def __init__(self, template: str, sname: str) -> None:
        self.doc = Document(template)
        self.name = sname
        self.setStyles()
        
    def setStyles(self):
        font = self.doc.styles['Normal'].font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        styles = self.doc.styles
        tcStyle = styles.add_style('Tc', WD_STYLE_TYPE.PARAGRAPH)
        tcStyle.base_style = styles['Normal']

    def numberRowsDataChecker(self, row: int,  header: bool, data: list[list[int|str]])-> bool:
        return row+int(header)== len(data)

    def numberDataItemsChecker(self, col: int, merge: bool, data: list[list[int|str]])-> bool:
        check = True
        # for line in data:
        # todo
        return check

    def creatAndFillTable(self,   header: list[str]|tuple[str],data: list[list[int|str]],col: int,row: int=1):
        table = self.doc.add_table(row, col, 'Table Grid')
        columns = table.columns
        columns[0].width = Cm(1.51)
        columns[1].width = Cm(2.31)
        columns[2].width = Cm(6.83)
        columns[3].width = Cm(6.83)
        hdrCells = table.rows[0].cells
        for i in range(col):
            hdrCells[i].text = header[i]
        for rec in data:
            rowCells = table.add_row().cells
            for i,cell in enumerate(rowCells):
                if isinstance(rec[i], Pic):
                    cell.paragraphs[0].add_run().add_picture(rec[i].path)
                else:
                    cell.text = str(rec[i])
        return table

    def pntSumaryTable(self,   data: list[list[int|str]]):
        table = self.doc.add_table(1, 4, 'Table Grid')
        columns = table.columns
        columns[0].width = Cm(1.0)
        columns[1].width = Cm(5.0)
        columns[2].width = Cm(6.49)
        columns[3].width = Cm(4.01)
        header = ['№ пп', 'Муниципальное образование', 'Проект перераспределения земель', 'Результат сравнения']
        hdrCells = table.rows[0].cells
        for i in range(4):
            hdrCells[i].text = header[i]
        for rec in data:
            rowCells = table.add_row().cells
            for i,cell in enumerate(rowCells):
                cell.text = str(rec[i])
        return table

    def distSumaryTable(self,   data: list[list[int|str]]):
        table = self.doc.add_table(1, 4, 'Table Grid')
        columns = table.columns
        columns[0].width = Cm(1.0)
        columns[1].width = Cm(3.94)
        columns[2].width = Cm(4.24)
        columns[3].width = Cm(7.31)
        header = ['№ пп', 'Муниципальное образование', 'Проект перераспределения земель', 'Результат сравнения']
        hdrCells = table.rows[0].cells
        for i in range(4):
            hdrCells[i].text = header[i]
        for rec in data:
            rowCells = table.add_row().cells
            for i,cell in enumerate(rowCells):
                cell.text = str(rec[i])
        return table
    
    def pntTable(self,   data: list[list[int|str]]):
        table = self.doc.add_table(1, 4, 'Table Grid')
        columns = table.columns
        columns[0].width = Cm(1.51)
        columns[1].width = Cm(2.31)
        columns[2].width = Cm(6.83)
        columns[3].width = Cm(6.83)
        header = ['Номер точки', 'Результат сравнения', 'Точка на архивном документе ГФДЗ', 'Точка на спутниковом общедоступном снимке']
        hdrCells = table.rows[0].cells
        for i in range(4):
            hdrCells[i].text = header[i]
        for rec in data:
            rowCells = table.add_row().cells
            for i,cell in enumerate(rowCells):
                if i<2:
                    cell.text = str(rec[i])
                else:
                    cell.paragraphs[0].add_run().add_picture(rec[i])
        return table

    def distTable(self,   data: list[list[int|str]]):
        table = self.doc.add_table(1, 2, 'Table Grid')
        columns = table.columns
        columns[0].width = Cm(4.0)
        columns[1].width = Cm(12.83)
        header = ['Номер точки', 'Расхождения, м']
        hdrCells = table.rows[0].cells
        for i in range(2):
            hdrCells[i].text = header[i]
        for rec in data:
            rowCells = table.add_row().cells
            for i,cell in enumerate(rowCells):
                cell.text = str(rec[i])
        return table
    def save(self):
        self.doc.save(self.name)


if __name__ == "__main__":
    import os
    from my_data_types import Pic
    url=os.path.join(os.path.dirname(__file__), 'test.docx')
    doc = MyDocx(url)
    header =('Номер точки','Результат сравнения','Картинка')
    records = (
        (1,'Результат сравнения', Pic('/home/archlex/.local/share/QGIS/QGIS3/profiles/default/python/plugins/binding_checker/test1.png')),
        (2,'Не совпадает', Pic('/home/archlex/.local/share/QGIS/QGIS3/profiles/default/python/plugins/binding_checker/test1.png'))
    )
    doc.creatAndFillTable(header, records, col=3,row=1)
    doc.save()
    

